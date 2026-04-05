import os
import config
import pymupdf.layout
import pymupdf4llm
from pathlib import Path
import glob
import re

os.environ["TOKENIZERS_PARALLELISM"] = "false"

def pdf_to_markdown(pdf_path, output_dir):
    doc = pymupdf.open(pdf_path)
    md = pymupdf4llm.to_markdown(doc,header=False, footer=False, page_separators=True, ignore_images=True, write_images=False, image_path=None)
    md = re.sub(r"\*\*==> picture .*? intentionally omitted <==\*\*", "", md)
    md = re.sub(r"--- end of page\.page_number=\d+ ---", "", md)
    md = re.sub(r"\n{3,}", "\n\n", md)
    md_cleaned = md.encode('utf-8', errors='surrogatepass').decode('utf-8', errors='ignore')
    output_path = Path(output_dir) / Path(doc.name).stem
    Path(output_path).with_suffix(".md").write_bytes(md_cleaned.encode('utf-8'))

def docx_to_markdown(docx_path, output_dir):
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(docx_path)
    lines = []
    HEADING_MAP = {
        "Heading 1": "#",
        "Heading 2": "##",
        "Heading 3": "###",
        "Heading 4": "####",
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
 
        style = para.style.name if para.style else "Normal"
 
        if style in HEADING_MAP:
            lines.append(f"{HEADING_MAP[style]} {text}")
        elif style.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    
    md = "\n\n".join(lines)
    md = re.sub(r"\n{3,}", "\n\n", md)
    md_cleaned = md.encode("utf-8", errors="surrogatepass").decode("utf-8", errors="ignore")
 
    output_path = Path(output_dir) / Path(docx_path).stem
    Path(output_path).with_suffix(".md").write_bytes(md_cleaned.encode("utf-8"))

def pdfs_to_markdowns(path_pattern, overwrite: bool = False):
    output_dir = Path(config.MARKDOWN_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)

    for pdf_path in map(Path, glob.glob(path_pattern)):
        md_path = (output_dir / pdf_path.stem).with_suffix(".md")
        if overwrite or not md_path.exists():
            pdf_to_markdown(pdf_path, output_dir)

def main():
    pdf_to_markdown(r"F:\project3\windowsMovieMaker.pdf","test")

if __name__ == "__main__":
    main()